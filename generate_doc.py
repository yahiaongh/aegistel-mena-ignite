import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Style Helpers
    NAVY = RGBColor(10, 25, 47)
    CYAN = RGBColor(0, 150, 214)
    DARK_GRAY = RGBColor(51, 51, 51)

    def set_cell_background(cell, hex_color):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_padding(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def add_heading(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = NAVY

    # Header Banner
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_cell = title_table.cell(0, 0)
    set_cell_background(title_cell, "0A192F")
    set_cell_padding(title_cell, top=200, bottom=200, left=200, right=200)

    p0 = title_cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run("AEGISTEL: AUTONOMOUS TELCO-AWARE AI GUARD")
    run0.font.name = "Arial"
    run0.font.size = Pt(20)
    run0.font.bold = True
    run0.font.color.rgb = RGBColor(255, 255, 255)

    p_sub = title_cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("GSMA MENA Ignite Hackathon 2026 · Idea Capture Template")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = CYAN

    doc.add_paragraph()

    # SECTION 1: METADATA & SUBMITTER DETAILS
    add_heading("1. Idea & Submitter Details")
    
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Idea Name:", "AegisTel - Autonomous Telco-Aware AI Guard Engine"),
        ("Submission Date:", "July 2026"),
        ("Submitter Name:", "Yahia Abdeldjalil"),
        ("Team Name & Contact:", "AegisTel Labs | yahia@aegistel.ai"),
        ("Selected Theme:", "Theme 7: Open Innovation (Spans Themes 2, 3, and 4)"),
        ("GSMA Pillar Alignment:", "GSMA Open Gateway Programmable Networks (3GPP TS 29.522)")
    ]
    for idx, (label, val) in enumerate(meta_data):
        c0, c1 = meta_table.cell(idx, 0), meta_table.cell(idx, 1)
        set_cell_background(c0, "F4F6F8")
        p0, p1 = c0.paragraphs[0], c1.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = NAVY
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = DARK_GRAY

    # SECTION 2: THEME SELECTION & RATIONALE
    add_heading("2. Theme Alignment Rationale")
    p = doc.add_paragraph()
    p.add_run("Selected Theme: ").bold = True
    p.add_run("Theme 7 — Open Innovation\n").font.color.rgb = CYAN
    p.add_run("Rationale: ").bold = True
    p.add_run(
        "AegisTel is designed as an autonomous cross-domain AI guard engine that intentionally spans multiple hackathon pillars:\n"
        "• Theme 4 (Secure Fintech & Anti-Fraud): Eliminates SIM swap attacks and validates line ownership prior to high-value wire transfers.\n"
        "• Theme 2 (Smart Cities & Urban Safety): Monitors real-time cell tower density and congestion in dense urban centers.\n"
        "• Theme 3 (Tourism & Pilgrimage): Automatically provisions dynamic 5G Quality-on-Demand (QoD) slices for emergency services during mega-events or pilgrimage crowds in Makkah."
    )

    # SECTION 3: IDEA SUMMARY
    add_heading("3. Idea Summary")
    p = doc.add_paragraph()
    p.add_run("• Problem Statement: ").bold = True
    p.add_run("Enterprise security and public safety applications suffer from fragmented telecom signals. Security workflows rely on vulnerable out-of-band OTPs susceptible to SIM swapping, while emergency teams lack automated, real-time access to 5G network priority slicing.\n")
    p.add_run("• Proposed Solution: ").bold = True
    p.add_run("AegisTel is an autonomous AI guard engine powered by Groq Llama-3.3-70B that ingests high-level application events, reasons about context, and dynamically chains 6 GSMA CAMARA Open Gateway APIs via Nokia Network-as-Code (NaC).\n")
    p.add_run("• Expected Benefits: ").bold = True
    p.add_run("Zero friction for users, automated fraud blocking before financial execution, real-time 5G QoS provisioning for emergency dispatches, and reduced operational costs for telecommunication service providers.")

    # SECTION 4: PROJECT TYPE & GSMA ALIGNMENT
    add_heading("4. Project Type & GSMA Alignment")
    p = doc.add_paragraph()
    p.add_run("Project Type: ").bold = True
    p.add_run("Autonomous Multi-Agent Telco Orchestrator (Middleware / AI Guard)\n")
    p.add_run("GSMA Open Gateway Alignment: ").bold = True
    p.add_run("Full alignment with standard CAMARA API specs (3GPP TS 29.522). AegisTel bridges raw 5G SBA control planes with application developers through zero-code autonomous tool invocation.")

    # SECTION 5: API USAGE
    add_heading("5. CAMARA API Usage Synopsis (6 APIs)")
    table = doc.add_table(rows=7, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["CAMARA API", "Technical Function", "Autonomous Action Trigger"]
    for i, h_text in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_background(cell, "0A192F")
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)

    api_data = [
        ("SIM Swap Detection", "Validates recent SIM changes", "Flags account takeover risk prior to transaction approval"),
        ("Location Verification", "Geofence circle verification", "Confirms physical device presence at ATM or incident area"),
        ("Quality on Demand (QoD)", "Provisions 5QI network slices", "Guarantees low latency / bandwidth for emergency feeds"),
        ("Number Verification", "Silent carrier authentication", "Eliminates SMS OTP reliance with header verification"),
        ("Congestion Insights", "Cell tower traffic density check", "Detects crowd gathering in high-density smart cities"),
        ("Device Reachability", "Connectivity & roaming status", "Verifies device is reachable before issuing QoD or dispatch")
    ]
    for row_idx, data in enumerate(api_data, start=1):
        bg = "F4F6F8" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = table.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8.5)

    # SECTION 6: WHERE 5G CONNECTIVITY FITS
    add_heading("6. Where 5G Connectivity Fits")
    fives_g_points = [
        ("5G Network Exposure Function (NEF): ", "Exposes 3GPP Northbound APIs (TS 29.522), allowing AegisTel to program network behavior over Nokia NaC."),
        ("5G Network Slicing & 5QI Customization: ", "Binds active traffic to dynamic 5G QoS Identifiers (5QI profiles) managed by the 5G Policy Control Function (PCF)."),
        ("5G Core Control Plane Integration: ", "Interrogates Unified Data Management (UDM) and Gateway Mobile Location Center (GMLC) without requiring client SDKs."),
        ("Multi-Access Edge Computing (MEC): ", "Offloads local video streams to 5G Edge Nodes during network congestion.")
    ]
    for title, desc in fives_g_points:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(title)
        r1.bold = True
        r1.font.color.rgb = NAVY
        r2 = p.add_run(desc)
        r2.font.color.rgb = DARK_GRAY

    doc.save("AegisTel_Idea_Capture_Template.docx")
    print("Successfully generated AegisTel_Idea_Capture_Template.docx")

if __name__ == "__main__":
    create_document()